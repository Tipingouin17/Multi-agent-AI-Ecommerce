"""
Carrier Contract Parsing Service

This service handles the upload and parsing of carrier contracts to extract:
- Pricing tables and rate cards
- Postal code zones
- Fuel surcharges
- Service level definitions
- Weight/dimension brackets
- Special handling fees

The service uses AI (OpenAI GPT) to intelligently parse various document formats
including PDF, Excel, CSV, and Word documents.

Technologies:
- OpenAI GPT-4 for intelligent document parsing
- PyPDF2 for PDF text extraction
- pandas for Excel/CSV processing
- python-docx for Word document processing
- PostgreSQL for storing parsed data
"""

import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any
import re

# Document processing
import PyPDF2
import pandas as pd
from docx import Document as DocxDocument

# AI processing
from openai import OpenAI

# Database
import psycopg2
from psycopg2.extras import RealDictCursor

# Configuration
DATABASE_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'port': os.getenv('DB_PORT', '5432'),
    'database': os.getenv('DB_NAME', 'multi_agent_ecommerce'),
    'user': os.getenv('DB_USER', 'postgres'),
    'password': os.getenv('DB_PASSWORD', 'postgres')
}

UPLOAD_DIR = os.getenv('CONTRACT_UPLOAD_DIR', '/app/storage/contracts')

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class ContractParsingService:
    """Service for parsing carrier contracts and extracting pricing information"""
    
    def __init__(self):
        self.db_conn = None
        self.openai_client = None
        self.initialize()
    
    def initialize(self):
        """Initialize database and OpenAI connections"""
        try:
            # Database connection
            self.db_conn = psycopg2.connect(**DATABASE_CONFIG)
            logger.info("Connected to database")
            
            # OpenAI client (API key from environment)
            self.openai_client = OpenAI()
            logger.info("OpenAI client initialized")
            
            # Ensure upload directory exists
            os.makedirs(UPLOAD_DIR, exist_ok=True)
            
        except Exception as e:
            logger.error(f"Initialization error: {str(e)}")
            raise
    
    def parse_contract(self, file_path: str, carrier_id: int, contract_metadata: Dict) -> Dict[str, Any]:
        """
        Parse a carrier contract and extract pricing information
        
        Args:
            file_path: Path to the contract file
            carrier_id: ID of the carrier
            contract_metadata: Additional metadata (contract_number, effective_date, etc.)
        
        Returns:
            Dictionary with parsed data and status
        """
        try:
            logger.info(f"Parsing contract: {file_path}")
            
            # Determine file type
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Extract text based on file type
            if file_ext == '.pdf':
                text_content = self.extract_pdf_text(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                text_content = self.extract_excel_text(file_path)
            elif file_ext == '.csv':
                text_content = self.extract_csv_text(file_path)
            elif file_ext == '.docx':
                text_content = self.extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Use AI to parse the contract
            parsed_data = self.ai_parse_contract(text_content, carrier_id)
            
            # Save contract record
            contract_id = self.save_contract_record(
                carrier_id=carrier_id,
                file_path=file_path,
                contract_metadata=contract_metadata,
                parsed_data=parsed_data
            )
            
            # Update rate cards
            self.update_rate_cards(carrier_id, contract_id, parsed_data)
            
            logger.info(f"Contract parsed successfully: {contract_id}")
            
            return {
                'success': True,
                'contract_id': contract_id,
                'parsed_data': parsed_data,
                'message': 'Contract parsed and rate cards updated successfully'
            }
            
        except Exception as e:
            logger.error(f"Error parsing contract: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    def extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += df.to_string() + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting Excel text: {str(e)}")
            raise
    
    def extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"Error extracting CSV text: {str(e)}")
            raise
    
    def extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def ai_parse_contract(self, text_content: str, carrier_id: int) -> Dict[str, Any]:
        """
        Use OpenAI GPT to parse contract text and extract structured data
        
        Args:
            text_content: Extracted text from contract
            carrier_id: ID of the carrier
        
        Returns:
            Dictionary with structured pricing data
        """
        try:
            # Create prompt for GPT
            prompt = f"""
You are an AI assistant specialized in parsing carrier shipping contracts. 
Analyze the following contract text and extract structured pricing information.

Extract the following information:
1. Base rates by weight brackets (e.g., 0-1kg, 1-5kg, 5-10kg, etc.)
2. Postal code zones and their definitions
3. Zone-based pricing (rates per zone)
4. Fuel surcharge percentage or formula
5. Service levels (standard, express, overnight, etc.) and their pricing
6. Dimensional weight rules
7. Special handling fees (oversized, fragile, dangerous goods, etc.)
8. Minimum charges
9. Residential delivery surcharges
10. Any other relevant pricing information

Format the output as a JSON object with the following structure:
{{
    "base_rates": [
        {{"weight_min": 0, "weight_max": 1, "zone": "1", "price": 5.50}},
        ...
    ],
    "postal_zones": [
        {{"zone_id": "1", "postal_codes": ["75001", "75002", ...], "description": "Paris"}},
        ...
    ],
    "fuel_surcharge": {{"type": "percentage", "value": 15.5}},
    "service_levels": [
        {{"code": "STANDARD", "name": "Standard Delivery", "price_modifier": 0}},
        {{"code": "EXPRESS", "name": "Express Delivery", "price_modifier": 10.00}},
        ...
    ],
    "dimensional_weight_divisor": 5000,
    "special_fees": [
        {{"type": "oversized", "threshold": "120cm", "fee": 25.00}},
        {{"type": "residential", "fee": 3.50}},
        ...
    ],
    "minimum_charge": 5.00
}}

Contract text:
{text_content[:15000]}  # Limit to first 15000 characters to avoid token limits

Provide only the JSON output, no additional text.
"""
            
            # Call OpenAI API
            response = self.openai_client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are a specialized AI for parsing shipping carrier contracts and extracting pricing data. Always respond with valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,  # Low temperature for consistent parsing
                max_tokens=4000
            )
            
            # Extract and parse JSON response
            response_text = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            if response_text.startswith('```'):
                response_text = re.sub(r'^```json?\n', '', response_text)
                response_text = re.sub(r'\n```$', '', response_text)
            
            parsed_data = json.loads(response_text)
            
            logger.info("Contract parsed successfully by AI")
            return parsed_data
            
        except json.JSONDecodeError as e:
            logger.error(f"Error parsing AI response as JSON: {str(e)}")
            logger.error(f"Response text: {response_text}")
            raise
        except Exception as e:
            logger.error(f"Error in AI parsing: {str(e)}")
            raise
    
    def save_contract_record(self, carrier_id: int, file_path: str, 
                           contract_metadata: Dict, parsed_data: Dict) -> int:
        """Save contract record to database"""
        try:
            with self.db_conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO carrier_contracts 
                    (carrier_id, contract_number, effective_date, expiry_date, 
                     file_path, parsed_data, status, uploaded_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                """, (
                    carrier_id,
                    contract_metadata.get('contract_number'),
                    contract_metadata.get('effective_date'),
                    contract_metadata.get('expiry_date'),
                    file_path,
                    json.dumps(parsed_data),
                    'active',
                    datetime.now()
                ))
                contract_id = cursor.fetchone()[0]
                self.db_conn.commit()
                return contract_id
        except Exception as e:
            logger.error(f"Error saving contract record: {str(e)}")
            self.db_conn.rollback()
            raise
    
    def update_rate_cards(self, carrier_id: int, contract_id: int, parsed_data: Dict):
        """Update carrier rate cards based on parsed contract data"""
        try:
            with self.db_conn.cursor() as cursor:
                # Update base rates
                if 'base_rates' in parsed_data:
                    for rate in parsed_data['base_rates']:
                        cursor.execute("""
                            INSERT INTO carrier_rate_cards 
                            (carrier_id, contract_id, weight_min, weight_max, 
                             zone, price, effective_date)
                            VALUES (%s, %s, %s, %s, %s, %s, CURRENT_DATE)
                            ON CONFLICT (carrier_id, weight_min, weight_max, zone) 
                            DO UPDATE SET 
                                price = EXCLUDED.price,
                                contract_id = EXCLUDED.contract_id,
                                effective_date = EXCLUDED.effective_date,
                                updated_at = CURRENT_TIMESTAMP
                        """, (
                            carrier_id,
                            contract_id,
                            rate.get('weight_min'),
                            rate.get('weight_max'),
                            rate.get('zone'),
                            rate.get('price')
                        ))
                
                # Update postal zones
                if 'postal_zones' in parsed_data:
                    for zone in parsed_data['postal_zones']:
                        cursor.execute("""
                            INSERT INTO carrier_postal_zones 
                            (carrier_id, zone_id, postal_codes, description)
                            VALUES (%s, %s, %s, %s)
                            ON CONFLICT (carrier_id, zone_id) 
                            DO UPDATE SET 
                                postal_codes = EXCLUDED.postal_codes,
                                description = EXCLUDED.description,
                                updated_at = CURRENT_TIMESTAMP
                        """, (
                            carrier_id,
                            zone.get('zone_id'),
                            zone.get('postal_codes'),
                            zone.get('description')
                        ))
                
                # Update fuel surcharge
                if 'fuel_surcharge' in parsed_data:
                    cursor.execute("""
                        INSERT INTO carrier_surcharges 
                        (carrier_id, surcharge_type, value, effective_date)
                        VALUES (%s, 'fuel', %s, CURRENT_DATE)
                        ON CONFLICT (carrier_id, surcharge_type) 
                        DO UPDATE SET 
                            value = EXCLUDED.value,
                            effective_date = EXCLUDED.effective_date,
                            updated_at = CURRENT_TIMESTAMP
                    """, (
                        carrier_id,
                        parsed_data['fuel_surcharge'].get('value')
                    ))
                
                # Update service levels
                if 'service_levels' in parsed_data:
                    for service in parsed_data['service_levels']:
                        cursor.execute("""
                            INSERT INTO carrier_service_levels 
                            (carrier_id, service_code, service_name, price_modifier)
                            VALUES (%s, %s, %s, %s)
                            ON CONFLICT (carrier_id, service_code) 
                            DO UPDATE SET 
                                service_name = EXCLUDED.service_name,
                                price_modifier = EXCLUDED.price_modifier,
                                updated_at = CURRENT_TIMESTAMP
                        """, (
                            carrier_id,
                            service.get('code'),
                            service.get('name'),
                            service.get('price_modifier')
                        ))
                
                self.db_conn.commit()
                logger.info(f"Rate cards updated for carrier {carrier_id}")
                
        except Exception as e:
            logger.error(f"Error updating rate cards: {str(e)}")
            self.db_conn.rollback()
            raise
    
    def get_contract_history(self, carrier_id: int) -> List[Dict]:
        """Get contract history for a carrier"""
        try:
            with self.db_conn.cursor(cursor_factory=RealDictCursor) as cursor:
                cursor.execute("""
                    SELECT id, contract_number, effective_date, expiry_date,
                           file_path, status, uploaded_at
                    FROM carrier_contracts
                    WHERE carrier_id = %s
                    ORDER BY uploaded_at DESC
                """, (carrier_id,))
                return [dict(row) for row in cursor.fetchall()]
        except Exception as e:
            logger.error(f"Error fetching contract history: {str(e)}")
            return []
    
    def cleanup(self):
        """Cleanup resources"""
        if self.db_conn:
            self.db_conn.close()
        logger.info("Contract Parsing Service stopped")


if __name__ == '__main__':
    # Example usage
    service = ContractParsingService()
    
    # Example: Parse a contract
    result = service.parse_contract(
        file_path='/path/to/contract.pdf',
        carrier_id=1,
        contract_metadata={
            'contract_number': 'DPD-2025-001',
            'effective_date': '2025-01-01',
            'expiry_date': '2025-12-31'
        }
    )
    
    print(json.dumps(result, indent=2))
    
    service.cleanup()

